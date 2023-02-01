
import numpy as np
import math
import docx
from docx import Document
from docx.shared import RGBColor

from MPM import doc
from changeWordColor import changeWordColor1
from charToNdarray import charToNdarray


"""
    用SB19方法来嵌入信息
    输入: 1、word文档的RGB列表,仅取R列表即可，需要Ndarray类型
          2、文字内容,可以传入字符串类型,但算法中需要Ndarray类型
    输出: 1、改后的RGB列表,仅有R列表
"""

# doc=docx.Document('demo03.docx')
# ALGORITHM: SB19 方法
def SB19(word_original_rgbList, secret_msg, k=3):
    RChannel_array = np.array(word_original_rgbList)
    secret_string = charToNdarray(secret_msg)
    print("secret_string",secret_string)
    n = 3  # 此算法在R、G、B通道中嵌入
    k2 = k + 1
    moshu0 = 2 ** (k + 1)  # 模数的底数
    moshu1 = 2 ** k  # 模数的底数
    num_RChannel_groups = RChannel_array.size
    print("num_RChannel_groups",num_RChannel_groups)
    # -----------------------------------------------------------------------------------
    # 从待嵌入bit串数据中取出k个比特，作为一组
    moshu = k * k
    # 秘密信息分组
    # print("secret_string.size",secret_string.size)
    num_secret_groups = math.ceil(secret_string.size / k)
    secret_group = np.zeros((num_secret_groups, k))
    # print("num_secret_groups",num_secret_groups)
    # print("secret_group",secret_group)
    for i in range(0, num_secret_groups, 1):
        for j in range(0, k, 1):
            if (i * k + j < secret_string.size):
                secret_group[i, j] = secret_string[i * k + j]
                # print("secret_group[i, j]",secret_group[i, j])
    # print("secret_group",secret_group)
    # 一组R通道嵌入一组secret_group的信息，多了不能嵌入,最后一组pixel不用于嵌入以防止错误
    assert (num_RChannel_groups >= num_secret_groups)
    # 每一组secret_group计算得到一个d值，d为（2n+1）进制进制的一个数
    secret_d_array = np.zeros(num_secret_groups)
    # print("secret_d_array",secret_d_array)
    for i in range(0, num_secret_groups, 1):
        # d代表一个（2n+1）进制的一个数
        d = 0
        for j in range(0, k, 1):
            d += secret_group[i, j] * (2 ** j)  # 将secret视为低位在前,1000,数值为1
        secret_d_array[i] = d
    print("secret_d_array",secret_d_array)

    # -----------------------------------------------------------------------------------
    # 开始进行嵌入
    embedded_RChannel_group = RChannel_array.copy()
    print("embedded_RChannel_group", embedded_RChannel_group)
    RChannels_group = RChannel_array.copy()
    # print("RChannels_group",RChannels_group)
    for i in range(0, num_secret_groups):
        x = 0
        for x in range(-1 * math.floor(moshu / 2), math.floor(moshu / 2) + 1, 1):
            f = (RChannels_group[i] + x) % moshu
            # print("f",f)
            if int(f) == int(secret_d_array[i]):
                if RChannels_group[i] + x < 0:
                    embedded_RChannel_group[i] = RChannels_group[i] + x + moshu
                elif RChannels_group[i] + x > 255:
                    embedded_RChannel_group[i] = RChannels_group[i] + x - moshu
                else:
                    embedded_RChannel_group[i] = RChannels_group[i] + x

                break
        tmp1 = embedded_RChannel_group[i] % moshu
        # print("tmp1",tmp1)
        tmp2 = int(secret_d_array[i])
        # print("tmp2", tmp2)
        assert (tmp1 == tmp2)
    print("embedded_RChannel_group",embedded_RChannel_group)

    # embedded_RChannel_pixels =[]
    # for i in range(len(embedded_RChannel_group)):
    #     embedded_RChannel_pixels.append(hex(embedded_RChannel_group[i])[2:].upper())
    # print("embedded_RChannel_pixels",embedded_RChannel_pixels)
    # print("listPixels",listPixels)

    # -----------------------------------------------------------------------------------
    # 使用了多少RChannel来进行嵌入
    # num_RChannels_changed = num_secret_groups * n
    # -----------------------------------------------------------------------------------
    # 输出
    return len(secret_string),embedded_RChannel_group

# 假设R列表为[255,255,255,255,255,255],秘密信息为[1 1 1 0 0 1 1 0 1 0 0 0 1 0 0 0 1 0 0 1 0 0 0 1]
# list = [255,255,255,255,255,255]

#将像素值转换为16进制的RGB值

R_channel = changeWordColor1(doc)
print("R_channel",R_channel)
used,res = SB19(R_channel,"我",4)
# print("res",res)
# print(used,res)












# # 获取全段落
# paragraphs=doc.paragraphs
# rgbList=[]
# fontName=[]
# textList=[]
#
# for paragraph in paragraphs:
#     for run in paragraph.runs:
#         fontname=run.font.name
#         fontName.append(fontname)
#         rgb = str(run.font.color.rgb)
#         rgbList.append(rgb)
#         print("fontname",fontname)
#         print("run.font.color.rgb",run.font.color.rgb)
#         print("run",run)
#     for text in paragraph.text:
#         textList.append(text)
#     print("rgbList",rgbList)
#
# print("textList",textList)
# documentFortest=Document();
#
# p=documentFortest.add_paragraph()
#
#
#
# count =0;
# print("testForWrite")
# for a in textList:
#     print(a)
#     run=p.add_run(a)
#     run.font.name=u'{0}'.format(fontName[count])
#     run._element.rPr.rFonts.set(qn('w:eastAsia'),run.font.name)
#     rgb=[str(rgbList[count][i:i+2])for i in range(0, 6, 2)]
#     print(rgb)
#     run.font.color.rgb=RGBColor(int(rgb[0],16),int(rgb[1],16),int(rgb[2],16))
#     count+=1
# documentFortest.save(r'C:\Users\ASUS\Desktop\2.docx')
# 要解决的问题：怎么才能生成的文档是已经隐藏好信息的文档
#先将隐藏后的文本像素转化为16进制值，转化为RGB值，将原来的RGB值的R通道换成新的R通道，再生成文档,