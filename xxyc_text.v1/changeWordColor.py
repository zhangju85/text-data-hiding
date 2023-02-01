import docx
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn
# import binascii
import numpy as np
# np.set_printoptions(threshold=np.inf)
# from HidingAlgorithm import SB19
#from MPM import *
from valueHash import orignalDocxAbst
from MPM import MPM2022
oda = orignalDocxAbst("Purchasing Contract.docx")
abst = oda.generateAbst()
print("OriginalHash: " + abst)
b=bin(int(abst,16))[2:] # 我写的
print("hash值的二进制编码",b)

'''
    读取载体文字，R通道的像素值
'''
rgbList=[]
fontName=[]
textList=[]
r_channel = []

def changeWordColor1(doc):
    # print(doc)
    # print(len(doc.paragraph[5].runs))
    # print(doc.paragraphs[1].runs[0].text)


    # g_channel = []
    # b_channel = []
    paragraphs=doc.paragraphs
    for paragraph in paragraphs:
        # print(paragraph)
        for run in paragraph.runs:
            rgb = str(run.font.color.rgb)
            fontname = run.font.name
            #
            if (len(run.text) > 1):
                for textLen in range(len(run.text)):
                    fontName.append(fontname)
                    rgbList.append(rgb)
            else:
                rgbList.append(rgb)
                fontName.append(fontname)
            # print(run.font.color.rgb)
            # rgbList.append(rgb)
        print("rgbList", rgbList)
            # print(run)
        #     fontname = run.font.name
        #     fontName.append(fontname)
        # print("fontName",fontName)
        for text in paragraph.text:
            textList.append(text)

    # print("rgbList", rgbList)
    print("fontName", len(fontName))
    print("rgbList", len(rgbList))
    print("textList", len(textList))

    count =0;
    # print("testForWrite")
    for a in textList:
        # print(a)
        # run=p.add_run(a)
        # run.font.name = u'{0}'.format(fontName[count])
        rgb = [str(rgbList[count][i:i + 2])for i in range(0, 6, 2)]
        # print("rgb",rgb)
        # print(rgb[0])
        r_channel.append(rgb[0])
        count += 1
        # g_channel.append(rgb[1])
        # count += 1
        # b_channel.append(rgb[2])
        # count += 1
    # print("r_channel",r_channel)
    # print("g_channel", g_channel)
    # print("b_channel", b_channel)

    listPixels = []
    for i in range(len(r_channel)):
        listPixels.append(int(r_channel[i], 16))
    print("listPixels",listPixels)



    #print("载体图像的像素值为",listPixels)
    print("载体图像的像素值为",len(listPixels))
    return listPixels
# RGB三通道嵌入
def test_for_RGB(doc):
    rgbList_1 = []
    fontName_1 = []
    textList_1 = []
    r_channel_1= []

    g_channel_1 = []
    b_channel_1 = []
    paragraphs=doc.paragraphs
    for paragraph in paragraphs:
        for run in paragraph.runs:
            rgb = str(run.font.color.rgb)
            fontname_1 = run.font.name
            #
            if (len(run.text) > 1):
                for textLen in range(len(run.text)):
                    fontName_1.append(fontname_1)
                    rgbList_1.append(rgb)
            else:
                rgbList_1.append(rgb)
                fontName_1.append(fontname_1)
        print("rgbList", rgbList_1)
        for text in paragraph.text:
            textList_1.append(text)
    print("fontName_1", len(fontName_1))
    print("rgbList_1", len(rgbList_1))
    print("textList_1", len(textList_1))

    count =0;
    for a in textList_1:
        rgb = [str(rgbList_1[count][i:i + 2])for i in range(0, 6, 2)]
        r_channel_1.append(rgb[0])
        g_channel_1.append(rgb[1])
        b_channel_1.append(rgb[2])
        count += 1


    listPixels_1 = []
    for i in range(len(r_channel_1)):
        listPixels_1.append(int(r_channel_1[i], 16))
        listPixels_1.append(int(g_channel_1[i], 16))
        listPixels_1.append(int(b_channel_1[i], 16))
    print("listPixels_1",listPixels_1)



    #print("载体图像的像素值为",listPixels)
    print("载体图像的像素值_1为",len(listPixels_1))
    return rgbList_1,textList_1 ,fontName_1,listPixels_1
#changeWordColor(doc)


# used,res = SB19(R_channel,abst,4)

# # 获取全段落
# paragraphs=doc.paragraphs
# rgbList=[]
# fontName=[]
# textList=[]
#
# for paragraph in paragraphs:
#     for run in paragraph.runs:
#
#         fontname=run.font.name
#         fontName.append(fontname)
#         rgb = str(run.font.color.rgb)
#         rgbList.append(rgb)
#         # print("fontname",fontname)
#         # print("run.font.color.rgb",run.font.color.rgb)
#         # print("run",run)
#     for text in paragraph.text:
#         textList.append(text)
#     # print("rgbList",rgbList)
#
# # print("textList",textList)

###

# print("testForWrite")


# doc = docx.Document('Purchasing Contract.docx')
# R_channel = changeWordColor1(doc)
# used, res = MPM2022(R_channel, abst, 4)
# print("”res”", res)
#将三通道加密并且写入文件，
def test_for_RGB_embeding(doc):
    from binhash import biAbsthash, abst
    from RecoverAlgorithm import MPMrecover

    #返回每个字内容，字体，每个字rgb排列
    rgbList_1,textList_1 ,fontName_1,listPixels_1=test_for_RGB(doc)
    used, res, sec, m = MPM2022(listPixels_1, abst, 3)

    print("嵌入信息abst: " + abst)
    secret_string = biAbsthash(abst)
    print("嵌入信息二进制: " , secret_string)
    print("原始RGB16进制: " , listPixels_1,)
    print("原始RGB16进制长度: ",  len(listPixels_1))
    print("原始RGB10进制: ", end="")
    for rgb in rgbList_1 :
        print(int(rgb[0:2],16),int(rgb[2:4],16),int(rgb[4:6],16),end=' ')

    print("\n返回的RGB: ", res,len(res))
    print(" : ",  res.shape[:])
    print("嵌入信息二进制: ", secret_string)
    print("恢复的秘密信息：",MPMrecover(np.array(sec),len(secret_string)))

    documentFortest = Document();
    p = documentFortest.add_paragraph()

    count = 0;
    for a in textList_1:
        # print(a)
        run = p.add_run(a)
        run.font.name = u'{0}'.format(fontName_1[count])
        run._element.rPr.rFonts.set(qn('w:eastAsia'), run.font.name)
        #将返回的RGB


        #print("每个字的变化后的rgb", res[count*3],res[count*3+1],res[count*3+2])
        run.font.color.rgb = RGBColor(int(res[count*3]), int(res[count*3]+1), int(res[count*3+2]))
        count += 1
    documentFortest.save(r'../xxyc_text.v1/test_RGB.docx')

if __name__ == '__main__':
    doc = docx.Document('text exam.docx')
    test_for_RGB_embeding(doc)

#
#
#     doc=docx.Document('Purchasing Contract.docx')
#
#     R_channel = changeWordColor1(doc)
#     documentFortest=Document();
#     p=documentFortest.add_paragraph()
#     print("R_channel.size", len(R_channel))
#     # used, res = MPM2022(R_channel, abst, 3)
#     # 返回 秘密字符串长度、R通道嵌入展平成一组、嵌入组数、嵌入率；
#     used, res, sec,m = MPM2022(R_channel, abst, 3)
#
#
#     print("”res1", len(sec))
#     print("”res2", sec.size)
#     print("”res3", sec.shape)
#     print("”res”", sec)
#     count =0;
#     for a in textList:
#         # print(a)
#         run=p.add_run(a)
#         run.font.name=u'{0}'.format(fontName[count])
#         run._element.rPr.rFonts.set(qn('w:eastAsia'),run.font.name)
#         rgb = [str(rgbList[count][i:i + 2]) for i in range(0, 6, 2)]
#         # doc = docx.Document('Purchasing Contract.docx')
#         # R_channel = changeWordColor1(doc)
#
#         print("res[count]",res[count])
#         # print("sec[count]", sec[count])
#         run.font.color.rgb=RGBColor(int(res[count]),int(rgb[1],16),int(rgb[2],16))
#         # print("run.font.color.rgb",run.font.color.rgb)
#         count+=1
#     documentFortest.save(r'../xxyc_text.v1/3.docx')
# # 要解决的问题：怎么才能生成的文档是已经隐藏好信息的文档
# #先将隐藏后的文本像素转化为16进制值，转化为RGB值，再生成文档
# ####



